from docx import Document
from docx.shared import Inches, Pt
import sleeperFunctions
import consts


def title_page(doc):
    doc.add_heading(sleeperFunctions.league_name(), 0)
    doc.add_picture('images/bigL.png', width=Inches(1.25))
    doc.add_page_break()


def week_results(doc):
    doc.add_heading("Week " + str(consts.WEEK()) + " Results", 1)
    doc.add_heading("Matchups", 2)
    results = sleeperFunctions.matchup_results(sleeperFunctions.matchups(consts.WEEK()))
    for i in range(0, len(results)):
        if i % 3 == 0 and i != 0:
            doc.add_paragraph("")
        elif i == 0:
            pass
        else:
            doc.add_paragraph(results[i])
    doc.add_heading("League Matchup Record", 2)
    lr = sleeperFunctions.league_results(results)
    for result in lr:
        doc.add_paragraph(str(result[0]) + ": " + str(result[1]) + "-" + str(result[2]) + "-" + str(result[3]))
    doc.add_heading("Roster Efficiency", 2)
    best_rosters = sleeperFunctions.team_efficiencies()
    for roster in best_rosters:
        doc.add_paragraph(sleeperFunctions.team_name(roster[0]) + " (" + sleeperFunctions.roster_id_to_owner(roster[0])
                          + "): " + "%.2f" % (100 * roster[1] / roster[2]) + "% - [" + "%.2f" % roster[1] + "/"
                          + "%.2f" % roster[2] + "]")
    doc.add_paragraph("")
    image_paragraph = doc.add_paragraph()
    for i in range(0, 24):
        run = image_paragraph.add_run()
        run.add_picture('images/gunning.png', width=Inches(0.25))


def season_results(doc):
    doc.add_heading("Season-Long Metrics", 1)
    doc.add_heading("League Matchup Record", 2)
    weekly_records = []
    for i in range(1, consts.WEEK() + 1):
        weekly_records.append(
            sleeperFunctions.league_results(sleeperFunctions.matchup_results(sleeperFunctions.matchups(i))))
    overall_matchups = sleeperFunctions.combine_records(weekly_records)
    for result in overall_matchups:
        doc.add_paragraph(str(result[0]) + ": " + str(result[1]) + "-" + str(result[2]) + "-" + str(result[3]))


def create_docx(filename):
    doc = Document()
    style = doc.styles['Normal']
    style.paragraph_format.space_after = Pt(5)

    # Title page
    title_page(doc)
    # Week results
    week_results(doc)
    # To-date Results
    season_results(doc)

    doc.save(filename)
