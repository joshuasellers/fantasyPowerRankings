import json
import requests
import consts
import glob, os
from operator import itemgetter
from docx import Document
from docx.shared import Inches, Pt
from docx2pdf import convert
import time


def user_name(userid):
    response = requests.get("https://api.sleeper.app/v1/user/" + userid)
    json_response = json.loads(response.text)
    return consts.OWNER_NAMES()[json_response['username']]


def league_owners():
    response = requests.get("https://api.sleeper.app/v1/league/" + consts.LEAGUE_ID() + "/rosters")
    json_response = json.loads(response.text)
    owners = []
    for team in json_response:
        owners.append(user_name(team['owner_id']))
    return ', '.join(owners)


def league_name():
    response = requests.get("https://api.sleeper.app/v1/league/" + consts.LEAGUE_ID())
    json_response = json.loads(response.text)
    return json_response['name']


def roster_id_to_owner(id):
    response = requests.get("https://api.sleeper.app/v1/league/" + consts.LEAGUE_ID() + "/rosters")
    json_response = json.loads(response.text)
    for team in json_response:
        if team['roster_id'] == id:
            return user_name(team['owner_id'])


def matchups(week):
    response = requests.get("https://api.sleeper.app/v1/league/" + consts.LEAGUE_ID() + "/matchups/" + str(week))
    json_response = json.loads(response.text)
    matchups = []
    for matchup in json_response:
        for opponent in json_response:
            if matchup['matchup_id'] == opponent['matchup_id'] and matchup['roster_id'] != opponent['roster_id']:
                matchups.append([matchup, opponent])
                json_response.remove(opponent)
    return matchups


def team_name(id):
    r = requests.get("https://api.sleeper.app/v1/league/" + consts.LEAGUE_ID() + "/rosters")
    jr = json.loads(r.text)
    uid = ""
    for team in jr:
        if team['roster_id'] == id:
            uid = team['owner_id']
    response = requests.get("https://api.sleeper.app/v1/league/" + consts.LEAGUE_ID() + "/users")
    json_response = json.loads(response.text)
    for team in json_response:
        if team['user_id'] == uid:
            if 'team_name' in team['metadata']:
                # if emoji.emoji_count(team['metadata']['team_name']) > 0:
                # return emoji.replace_emoji(team['metadata']['team_name'], "")
                return team['metadata']['team_name']
            else:
                return "Team " + team['display_name']


def matchup_results(matchups):
    mr = []
    for i in range(0, len(matchups)):
        mr.append("Matchup " + str(i+1))
        if sum(matchups[i][0]['starters_points']) >= sum(matchups[i][1]['starters_points']):
            mr.append(team_name(matchups[i][0]['roster_id']) + " ("
                      + roster_id_to_owner(matchups[i][0]['roster_id'])
                      + "): " + str("%.2f" % sum(matchups[i][0]['starters_points'])))
            mr.append(team_name(matchups[i][1]['roster_id']) + " ("
                      + roster_id_to_owner(matchups[i][1]['roster_id'])
                      + "): " + str("%.2f" % sum(matchups[i][1]['starters_points'])))
        else:
            mr.append(team_name(matchups[i][1]['roster_id']) + " ("
                      + roster_id_to_owner(matchups[i][1]['roster_id'])
                      + "): " + str("%.2f" % sum(matchups[i][1]['starters_points'])))
            mr.append(team_name(matchups[i][0]['roster_id']) + " ("
                      + roster_id_to_owner(matchups[i][0]['roster_id'])
                      + "): " + str("%.2f" % sum(matchups[i][0]['starters_points'])))
    return mr


def league_results(results):
    lr = []
    for i in range(0, len(results)):
        if i % 3 != 0:
            record = [results[i].split(":")[0], 0, 0, 0]
            score = float(results[i].split(":")[1])
            for j in range(0,len(results)):
                if i != j and j % 3 != 0:
                    if score > float(results[j].split(":")[1]):
                        record[1] = record[1] + 1
                    elif score < float(results[j].split(":")[1]):
                        record[2] = record[2] + 1
                    else:
                        record[3] = record[3] + 1
            lr.append(record)
    lr = sorted(lr, key=itemgetter(1), reverse=True)
    return lr


def combine_records(weekly_records):
    records = {}
    for week in weekly_records:
        for record in week:
            if record[0] in records:
                records[record[0]][0] += record[1]
                records[record[0]][1] += record[2]
                records[record[0]][2] += record[3]
            else:
                records[record[0]] = [record[1], record[2], record[3]]
    unsorted_records = []
    for owner in records:
        unsorted_records.append([owner] + records[owner])
    return sorted(unsorted_records, key=itemgetter(1), reverse=True)


def create_docx(filename):
    doc = Document()
    style = doc.styles['Normal']
    style.paragraph_format.space_after = Pt(5)

    # Title page
    doc.add_heading(league_name(), 0)
    doc.add_picture('images/bigL.png', width=Inches(1.25))
    doc.add_page_break()

    # Week results
    doc.add_heading("Week " + str(consts.WEEK()) + " Results", 1)
    doc.add_heading("Matchups", 2)
    results = matchup_results(matchups(consts.WEEK()))
    for i in range(0, len(results)):
        if i % 3 == 0 and i != 0:
            doc.add_paragraph("")
        elif i == 0:
            pass
        else:
            doc.add_paragraph(results[i])
    doc.add_heading("League Matchup Record", 2)
    lr = league_results(results)
    for result in lr:
        doc.add_paragraph(str(result[0]) + ": " + str(result[1]) + "-" + str(result[2]) + "-" + str(result[3]))
    doc.add_heading("Roster Efficiency", 2)
    best_rosters = team_efficiencies()
    for roster in best_rosters:
        doc.add_paragraph(team_name(roster[0]) + " (" + roster_id_to_owner(roster[0])
                          + "): " + "%.2f" % (100 * roster[1]/roster[2]) + "% - [" + "%.2f" % roster[1] + "/"
                          + "%.2f" % roster[2] + "]")

    # To-date Results
    doc.add_heading("Season-Long Metrics", 1)
    doc.add_heading("League Matchup Record", 2)
    weekly_records = []
    for i in range(1, consts.WEEK()+1):
        weekly_records.append(league_results(matchup_results(matchups(i))))
    overall_matchups = combine_records(weekly_records)
    for result in overall_matchups:
        doc.add_paragraph(str(result[0]) + ": " + str(result[1]) + "-" + str(result[2]) + "-" + str(result[3]))

    doc.save(filename)


def update_player_data():
    response = requests.get("https://api.sleeper.app/v1/players/nfl")
    if os.path.exists('player_data/player_data.txt'):
        os.remove('player_data/player_data.txt')
    with open('player_data/player_data.txt', 'w') as f:
        f.write(response.text)


def get_player_data():
    if os.path.exists('player_data/player_data.txt'):
        with open('player_data/player_data.txt', 'r') as f:
            return json.load(f)


def team_efficiencies():
    data = get_player_data()
    league_response = requests.get(
        "https://api.sleeper.app/v1/league/" + consts.LEAGUE_ID())
    json_league_response = json.loads(league_response.text)
    starting_roster = list(filter(lambda x: x != 'BN', json_league_response['roster_positions']))
    matchups_response = requests.get(
        "https://api.sleeper.app/v1/league/" + consts.LEAGUE_ID() + "/matchups/" + str(consts.WEEK()))
    json_matchups_response = json.loads(matchups_response.text)
    best_rosters = []
    for matchup in json_matchups_response:
        optimal_roster = []
        optimal_score = 0
        for position in starting_roster:
            max_points = 0
            best_player = ''
            for player in matchup['players_points']:
                orf = [x[0] for x in optimal_roster if x[0] == player]
                if player in orf:
                    continue
                if data[player]['position'] == position:
                    if matchup['players_points'][player] >= max_points:
                        max_points = matchup['players_points'][player]
                        best_player = player
                elif position == 'SUPER_FLEX':
                    if matchup['players_points'][player] >= max_points:
                        max_points = matchup['players_points'][player]
                        best_player = player
                elif position == 'FLEX' and data[player]['position'] != 'QB':
                    if matchup['players_points'][player] >= max_points:
                        max_points = matchup['players_points'][player]
                        best_player = player
            optimal_roster.append([best_player, position, max_points, data[best_player]['full_name']])
            optimal_score += max_points
        best_rosters.append([matchup['roster_id'], matchup['points'], optimal_score, optimal_roster,
                             matchup['points']/optimal_score])
    return sorted(best_rosters, key=itemgetter(4), reverse=True)


if __name__ == '__main__':
    # update_player_data() # call every once in a while to keep this up to date
    filename = 'week' + str(consts.WEEK()) + 'results'
    print("Removing any previous files")
    for f in glob.glob("*.docx"):
        os.remove(f)
    for f in glob.glob("*.pdf"):
        os.remove(f)
    time.sleep(5)
    print("Creating new docx file")
    create_docx(filename + '.docx')
    time.sleep(2.5)
    print("Converting new docx file to a PDF file")
    if os.path.exists(filename + '.docx'):
        convert(filename + '.docx')
